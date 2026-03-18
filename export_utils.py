import streamlit as st
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document

def save_as_pdf(text):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 800, "Extracted Text:")
    
    y_position = 780
    for line in text.split("\n"):
        pdf.drawString(100, y_position, line)
        y_position -= 20  # Adjust line spacing

    pdf.save()
    buffer.seek(0)

    st.download_button(
        label="Download PDF",
        data=buffer,
        file_name="extracted_text.pdf",
        mime="application/pdf"
    )

def save_as_word(text):
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Word",
        data=buffer,
        file_name="extracted_text.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
